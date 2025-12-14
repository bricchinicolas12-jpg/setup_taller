# orden_docx.py
from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[^\w\s\-\.]", "", s, flags=re.UNICODE)
    s = re.sub(r"\s+", "_", s)
    return s[:120] if s else "orden"


def _fmt_date(val: Any) -> str:
    if not val:
        return ""
    # val puede venir como 'YYYY-MM-DD' o datetime/date
    if hasattr(val, "strftime"):
        return val.strftime("%Y-%m-%d")
    return str(val)[:10]


def _fmt_time(val: Any) -> str:
    if not val:
        return ""
    s = str(val)
    return s[:5]  # HH:MM


def _add_cell_border(cell):
    """Borde simple para celdas (mejora impresión)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')     # grosor
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '999999')
        tcBorders.append(element)

    tcPr.append(tcBorders)


def _set_default_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Márgenes imprimibles (A4)
    section = doc.sections[0]
    # docx usa pulgadas internamente; estas aprox son 2cm
    section.top_margin = Pt(57)
    section.bottom_margin = Pt(57)
    section.left_margin = Pt(57)
    section.right_margin = Pt(57)


def generar_docx_orden(
    orden: Dict[str, Any],
    output_dir: str,
    filename: Optional[str] = None,
) -> str:
    """
    orden: dict con campos de la orden (id, fecha, cliente, equipo, etc.)
    output_dir: carpeta donde guardar el .docx
    filename: opcional, nombre del archivo. Si no se pasa, se genera automático.
    Devuelve la ruta completa del archivo creado.
    """
    os.makedirs(output_dir, exist_ok=True)

    oid = orden.get("id") or orden.get("nro") or ""
    cliente = orden.get("nombre_contacto") or orden.get("cliente_nombre") or orden.get("cliente") or ""
    equipo = orden.get("equipo_texto") or orden.get("equipo") or orden.get("equipo_descripcion") or ""

    if not filename:
        base = f"Orden_{oid}_{_safe_filename(cliente)}_{_safe_filename(equipo)}"
        filename = base + ".docx"

    path = os.path.join(output_dir, filename)

    doc = Document()
    _set_default_styles(doc)

    # Encabezado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ORDEN DE SERVICIO")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # espacio

    # Tabla de datos principales
    t = doc.add_table(rows=0, cols=4)
    t.style = "Table Grid"

    def add_row(label1, value1, label2, value2):
        row = t.add_row().cells
        row[0].text = str(label1)
        row[1].text = str(value1 or "")
        row[2].text = str(label2)
        row[3].text = str(value2 or "")
        for c in row:
            _add_cell_border(c)

        # labels en negrita
        row[0].paragraphs[0].runs[0].bold = True
        row[2].paragraphs[0].runs[0].bold = True

    add_row("N°", oid, "Estado", orden.get("estado", ""))
    add_row("Fecha ingreso", _fmt_date(orden.get("fecha")), "Hora ingreso", _fmt_time(orden.get("hora_ingreso")))
    add_row("Cliente / Contacto", cliente, "Teléfono", orden.get("telefono_contacto") or orden.get("telefono") or "")
    add_row("Equipo", equipo, "S/N", orden.get("serie_texto") or orden.get("serie") or "")
    add_row("Fecha salida", _fmt_date(orden.get("fecha_salida")), "Hora salida", _fmt_time(orden.get("hora_salida")))
    add_row("Fecha regreso", _fmt_date(orden.get("fecha_regreso")), "Hora regreso", _fmt_time(orden.get("hora_regreso")))
    add_row("Importe", orden.get("importe", ""), "Accesorios", orden.get("accesorios", ""))

    doc.add_paragraph()  # espacio

    # Secciones de texto largo
    def section(title: str, content: Any):
        h = doc.add_paragraph()
        r = h.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        doc.add_paragraph(str(content or ""))

    section("Falla", orden.get("falla"))
    section("Reparación", orden.get("reparacion"))
    section("Repuestos", orden.get("repuestos"))
    section("Observaciones", orden.get("observaciones"))

    # Retiro (si existe)
    fr = _fmt_date(orden.get("fecha_retiro"))
    hr = _fmt_time(orden.get("hora_retiro"))
    if fr or hr:
        section("Retiro", f"{fr} {hr}".strip())

    doc.add_paragraph()
    doc.add_paragraph("Firma / Aclaración: ________________________________")
    doc.add_paragraph("DNI: ______________________   Fecha: ____/____/______")

    # Pie con timestamp de generación
    doc.add_paragraph()
    pie = doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pie.runs[0].italic = True
    pie.runs[0].font.size = Pt(9)

    doc.save(path)
    return path
